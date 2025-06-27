# utils/report_gen.py

import os
import cv2
from docx import Document
from docx.shared import Inches
from utils.gpt4all_model import generate_response  # يعتمد على GPT4All المحلي

def generate_llm_report(fault_type, confidence, timestamp, suppliers_text):
    prompt = f"""
An AI-based drone system detected a fault in a power line insulator.

- Fault Type: {fault_type}
- Confidence Level: {confidence * 100:.2f}%
- Timestamp: {timestamp}

Relevant suppliers that might help:
{suppliers_text}

Please describe the possible impact of this fault and recommend appropriate maintenance or action.
"""
    output = generate_response(prompt, max_tokens=300)
    return output.strip()

def save_report_to_word(report_text, filename="Report.docx", output_dir="reports", image_path=None, suppliers=None, detections=None):
    os.makedirs(output_dir, exist_ok=True)
    save_path = os.path.join(output_dir, filename)

    doc = Document()
    doc.add_heading("🔧 Insulator Fault Inspection Report", level=1)

    # صورة الكشف
    if image_path and os.path.exists(image_path):
        if detections:
            frame = cv2.imread(image_path)
            for det in detections:
                x1, y1, x2, y2 = map(int, det["bbox"])
                label = det["label"]
                cv2.rectangle(frame, (x1, y1), (x2, y2), (0, 0, 255), 2)
                cv2.putText(frame, label, (x1, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (0, 0, 255), 2)
            temp_img_path = os.path.join(output_dir, "temp_detected_image.jpg")
            cv2.imwrite(temp_img_path, frame)
            doc.add_picture(temp_img_path, width=Inches(5))
            os.remove(temp_img_path)
        else:
            doc.add_picture(image_path, width=Inches(5))

        doc.add_paragraph("")

    # نص التقرير
    doc.add_heading("📄 AI Report", level=2)
    doc.add_paragraph(report_text)

    # الموردين
    if suppliers:
        doc.add_heading("🏭 Suggested Suppliers", level=2)
        for s in suppliers:
            doc.add_paragraph(f"🔹 {s.get('name', '')}\n{s.get('description', '')}\n{s.get('url', '')}")

    doc.save(save_path)
    return save_path
